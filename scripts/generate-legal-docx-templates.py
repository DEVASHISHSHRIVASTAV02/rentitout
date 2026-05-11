from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
TEMPLATES_DIR = ROOT / "public" / "legal-templates"

SOURCE_TARGETS = [
    (
        TEMPLATES_DIR / "appliance-rental-agreement-comprehensive.txt",
        TEMPLATES_DIR / "appliance-rental-agreement-comprehensive.docx",
    ),
    (
        TEMPLATES_DIR / "appliance-rental-agreement-short-form.txt",
        TEMPLATES_DIR / "appliance-rental-agreement-short-form.docx",
    ),
    (
        TEMPLATES_DIR / "appliance-handover-checklist.txt",
        TEMPLATES_DIR / "appliance-handover-checklist.docx",
    ),
]


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    heading_1 = doc.styles["Heading 1"]
    heading_1.font.name = "Calibri"
    heading_1.font.size = Pt(16)

    heading_2 = doc.styles["Heading 2"]
    heading_2.font.name = "Calibri"
    heading_2.font.size = Pt(13)

    heading_3 = doc.styles["Heading 3"]
    heading_3.font.name = "Calibri"
    heading_3.font.size = Pt(11)


def is_caps_heading(line: str) -> bool:
    letters = [ch for ch in line if ch.isalpha()]
    if not letters:
        return False
    return all(ch.isupper() for ch in letters)


def looks_like_section_heading(line: str) -> bool:
    return bool(
        re.match(r"^\d+\.", line)
        or re.match(r"^\d+\)", line)
        or re.match(r"^[A-Z]\)", line)
        or line.startswith("SCHEDULE ")
    )


def add_content(doc: Document, raw_text: str) -> None:
    lines = raw_text.splitlines()
    if not lines:
        return

    # Title block
    title_lines = [line.strip() for line in lines[:3] if line.strip()]
    if title_lines:
        title = doc.add_paragraph(title_lines[0], style="Heading 1")
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if len(title_lines) > 1:
        subtitle = doc.add_paragraph(title_lines[1], style="Heading 2")
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if len(title_lines) > 2:
        meta = doc.add_paragraph(title_lines[2])
        meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("")

    for raw_line in lines[3:]:
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph("")
            continue

        if re.fullmatch(r"[-=]{5,}", stripped):
            continue

        if stripped == "IMPORTANT NOTICE":
            doc.add_paragraph(stripped, style="Heading 2")
            continue

        if looks_like_section_heading(stripped):
            doc.add_paragraph(stripped, style="Heading 2")
            continue

        if is_caps_heading(stripped) and len(stripped) <= 70:
            doc.add_paragraph(stripped, style="Heading 3")
            continue

        if stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
            continue

        if stripped.startswith("[ ]"):
            doc.add_paragraph(stripped, style="List Bullet")
            continue

        paragraph = doc.add_paragraph(stripped)
        if stripped.endswith(":"):
            run = paragraph.runs[0]
            run.bold = True


def generate_docx(source_path: Path, target_path: Path) -> None:
    raw_text = source_path.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc)
    add_content(doc, raw_text)
    doc.save(target_path)


def main() -> None:
    for source_path, target_path in SOURCE_TARGETS:
        generate_docx(source_path, target_path)
        print(f"Generated: {target_path.relative_to(ROOT)}")


if __name__ == "__main__":
    main()

